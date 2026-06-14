from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
import logging
import json
import uuid
import os
import shutil
import tempfile
from typing import List
from datetime import datetime
from app.schemas import ChatRequest
from app.client import azure_client_manager
from app.services import local_form_filler_agent
from app.db import (
    create_session,
    save_message,
    save_attachment,
    get_sessions,
    get_session_messages,
    delete_session
)

logger = logging.getLogger("app.routes")

router = APIRouter()

@router.get("/sessions")
def read_sessions():
    try:
        return get_sessions()
    except Exception as e:
        logger.error(f"Error fetching sessions: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/session/{session_id}")
def read_session_messages(session_id: str):
    try:
        messages = get_session_messages(session_id)
        # Format dates if needed or just return raw
        return {"session_id": session_id, "messages": messages}
    except Exception as e:
        logger.error(f"Error fetching session messages: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/session/{session_id}")
def remove_session(session_id: str):
    try:
        delete_session(session_id)
        
        # Delete session folder from uploads directory if it exists
        upload_dir = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "uploads", session_id))
        uploads_parent = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "uploads"))
        if os.path.commonpath([uploads_parent, upload_dir]) == uploads_parent and upload_dir != uploads_parent:
            if os.path.exists(upload_dir):
                shutil.rmtree(upload_dir, ignore_errors=True)
                logger.info(f"Deleted uploads folder for session {session_id}")
                
        return {"status": "success", "message": f"Session {session_id} deleted"}
    except Exception as e:
        logger.error(f"Error deleting session: {e}")
        raise HTTPException(status_code=500, detail=str(e))

_easyocr_reader = None

def get_easyocr_reader():
    global _easyocr_reader
    if _easyocr_reader is None:
        import easyocr
        logger.info("Initializing EasyOCR Reader (CPU mode)...")
        _easyocr_reader = easyocr.Reader(['en'], gpu=False)
    return _easyocr_reader

@router.post("/upload/{session_id}")
async def upload_documents(session_id: str, files: List[UploadFile] = File(...)):
    import asyncio
    
    async def upload_event_generator():
        try:
            # Create directory at backend/uploads/session-id/
            upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "uploads", session_id)
            os.makedirs(upload_dir, exist_ok=True)
            
            extracted_results = []
            all_progress_events = []
            
            for file in files:
                if not file.filename:
                    continue
                    
                file_path = os.path.join(upload_dir, file.filename)
                
                # Emit OCR start event
                ocr_start_ev = {'type': 'ocr_start', 'filename': file.filename}
                all_progress_events.append(ocr_start_ev)
                yield f"data: {json.dumps(ocr_start_ev)}\n\n"
                
                # Read and save file content
                content = await file.read()
                with open(file_path, "wb") as f:
                    f.write(content)
                    
                filesize = len(content)
                filetype = file.content_type or file.filename.split(".")[-1]
                
                extracted_text = ""
                IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".heic", ".heif", ".avif")
                is_image = file.filename.lower().endswith(IMAGE_EXTENSIONS)
                
                # Extract text based on file type
                if file.filename.lower().endswith(".pdf"):
                    try:
                        import fitz  # PyMuPDF
                        doc = fitz.open(file_path)
                        text_parts = []
                        for page_num in range(len(doc)):
                            page = doc.load_page(page_num)
                            page_text = page.get_text().strip()
                            if len(page_text) > 15:
                                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                            else:
                                # Scanned page OCR
                                try:
                                    pix = page.get_pixmap(dpi=150)
                                    img_data = pix.tobytes("png")
                                    from PIL import Image
                                    import io
                                    import numpy as np
                                    img = Image.open(io.BytesIO(img_data))
                                    if img.mode != 'RGB':
                                        img = img.convert('RGB')
                                    reader = get_easyocr_reader()
                                    result = reader.readtext(np.array(img), detail=0)
                                    ocr_text = " ".join(result).strip()
                                    if ocr_text:
                                        text_parts.append(f"--- Page {page_num + 1} (Scanned OCR) ---\n{ocr_text}")
                                    else:
                                        text_parts.append(f"--- Page {page_num + 1} (Scanned OCR - Empty) ---")
                                except Exception as page_ocr_err:
                                    logger.error(f"Failed to OCR scanned PDF page {page_num + 1}: {page_ocr_err}")
                                    text_parts.append(f"--- Page {page_num + 1} (OCR Error) ---")
                        extracted_text = "\n\n".join(text_parts).strip()
                    except Exception as pdf_err:
                        logger.error(f"Failed to extract text from PDF {file.filename}: {pdf_err}")
                        extracted_text = ""
                elif file.filename.lower().endswith(".docx"):
                    try:
                        import docx
                        doc = docx.Document(file_path)
                        text_parts = []
                        for para in doc.paragraphs:
                            if para.text.strip():
                                text_parts.append(para.text)
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                                if row_text:
                                    text_parts.append(" | ".join(row_text))
                        extracted_text = "\n".join(text_parts).strip()
                    except Exception as docx_err:
                        logger.error(f"Failed to extract text from DOCX {file.filename}: {docx_err}")
                        extracted_text = ""
                elif is_image:
                    try:
                        from PIL import Image
                        import numpy as np
                        img = Image.open(file_path)
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        reader = get_easyocr_reader()
                        result = reader.readtext(np.array(img), detail=0)
                        extracted_text = " ".join(result).strip()
                    except Exception as ocr_err:
                        logger.error(f"Failed to OCR image {file.filename}: {ocr_err}")
                        extracted_text = ""
                else:
                    # Text files
                    try:
                        extracted_text = content.decode("utf-8", errors="ignore").strip()
                    except Exception:
                        extracted_text = ""
                
                # Handle empty/passport photo logic
                if is_image and not extracted_text:
                    extracted_text = "Passport size photo"
                elif not extracted_text:
                    extracted_text = "No text extracted"
                    
                # Keep track of result
                item = {
                    "filename": file.filename,
                    "filesize": filesize,
                    "extracted_text": extracted_text,
                    "filetype": filetype
                }
                extracted_results.append(item)
                
                # Emit OCR end event
                preview = extracted_text[:80] + "..." if len(extracted_text) > 80 else extracted_text
                ocr_end_ev = {'type': 'ocr_end', 'filename': file.filename, 'text_preview': preview}
                all_progress_events.append(ocr_end_ev)
                yield f"data: {json.dumps(ocr_end_ev)}\n\n"
                
            # Create a user message in DB representing the upload
            upload_msg_id = str(uuid.uuid4())
            filenames_str = ", ".join([r["filename"] for r in extracted_results])
            save_message(
                upload_msg_id,
                session_id,
                "user",
                f"Uploaded attachments: {filenames_str}",
                datetime.utcnow().isoformat()
            )
            
            # Save raw attachment references in DB for user message
            for r in extracted_results:
                save_attachment(
                    upload_msg_id,
                    r["filename"],
                    r["filesize"],
                    f"/uploads/{session_id}/{r['filename']}",
                    extracted_text=r.get("extracted_text", "")
                )
                
            # 1. Fetch remote agent's response from SQLite DB history
            db_messages = get_session_messages(session_id)
            remote_response = ""
            for msg in reversed(db_messages):
                if msg["role"] == "assistant" and msg["content"] and "Hi, by looking" in msg["content"]:
                    remote_response = msg["content"]
                    break
            if not remote_response:
                for msg in reversed(db_messages):
                    if msg["role"] == "assistant" and msg["content"] and msg["id"] != upload_msg_id:
                        remote_response = msg["content"]
                        break
    
            # 2. Setup progress queue and run local document processing workflow
            queue = asyncio.Queue()
            loop = asyncio.get_running_loop()
            
            def on_progress(event_type: str, data: dict):
                event = {"type": event_type, **data}
                all_progress_events.append(event)
                loop.call_soon_threadsafe(queue.put_nowait, event)
                
            from app.services.agent_workflow import local_document_processing_workflow
            
            input_data = {
                "remote_response": remote_response,
                "files_list": extracted_results,
                "session_id": session_id,
                "on_progress": on_progress
            }
            
            # Start workflow execution in background task
            workflow_task = asyncio.create_task(local_document_processing_workflow.run(input_data))
            
            # Stream events as they are pushed to queue
            while not workflow_task.done() or not queue.empty():
                try:
                    event = await asyncio.wait_for(queue.get(), timeout=0.1)
                    yield f"data: {json.dumps(event)}\n\n"
                except asyncio.TimeoutError:
                    continue
                    
            # Await final result
            run_result = await workflow_task
            outputs = run_result.get_outputs()
            result_dict = outputs[-1] if outputs else {}
            
            if isinstance(result_dict, dict):
                final_text = result_dict.get("final_response", "Failed to process uploaded documents.")
                processed_files = result_dict.get("processed_files", [])
                citation_map = result_dict.get("citation_map", {})
            else:
                final_text = result_dict if isinstance(result_dict, str) else "Failed to process uploaded documents."
                processed_files = []
                citation_map = {}
                
            # 3. Save assistant response (local agent extraction) to DB
            assistant_msg_id = str(uuid.uuid4())
            save_message(
                assistant_msg_id,
                session_id,
                "assistant",
                final_text,
                datetime.utcnow().isoformat(),
                progress_events=json.dumps(all_progress_events),
                citation_map=json.dumps(citation_map) if citation_map else None
            )
            
            # 4. Save processed attachments in DB linked ONLY to assistant response
            for pf in processed_files:
                fname = pf.get("filename")
                if not fname:
                    continue
                fpath = os.path.join(upload_dir, fname)
                if os.path.exists(fpath) and os.path.isfile(fpath):
                    save_attachment(
                        assistant_msg_id,
                        fname,
                        os.path.getsize(fpath),
                        f"/uploads/{session_id}/{fname}",
                        extracted_text=pf.get("extracted_text", "")
                    )
            
            # Emit done event with payload
            yield f"data: {json.dumps({'type': 'done', 'final_response': final_text, 'processed_files': processed_files, 'citation_map': citation_map})}\n\n"
            
        except Exception as e:
            logger.error(f"Error handling file upload: {e}")
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
            
    return StreamingResponse(upload_event_generator(), media_type="text/event-stream")

@router.post("/local_chat")
async def handle_local_chat(payload: ChatRequest):
    import asyncio
    
    session_id = payload.session_id
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")

    last_user_message = ""
    if payload.messages:
        for m in reversed(payload.messages):
            if m.role == "user" and m.content:
                last_user_message = m.content
                break
    if not last_user_message:
        last_user_message = payload.message or ""

    # Save user message to database
    user_msg_id = str(uuid.uuid4())
    save_message(
        user_msg_id,
        session_id,
        "user",
        last_user_message,
        datetime.utcnow().isoformat()
    )

    # 1. Fetch remote agent's response and files list from DB history
    db_messages = get_session_messages(session_id)
    
    # Rebuild files list from the most recent assistant message's attachments
    files_list = []
    last_assistant_msg = None
    for msg in reversed(db_messages):
        # Skip user message we just saved and find the last assistant message
        if msg["role"] == "assistant" and msg.get("attachments"):
            last_assistant_msg = msg
            break

    if last_assistant_msg:
        for att in last_assistant_msg["attachments"]:
            files_list.append({
                "filename": att["name"],
                "filesize": att["size"] or 0,
                "filetype": att["name"].split(".")[-1],
                "extracted_text": att.get("extracted_text") or ""
            })
    else:
        # Fallback to user uploaded raw attachments if no assistant response has attachments yet
        for msg in db_messages:
            if msg["role"] == "user" and msg.get("attachments") and msg["id"] != user_msg_id:
                for att in msg["attachments"]:
                    files_list.append({
                        "filename": att["name"],
                        "filesize": att["size"] or 0,
                        "filetype": att["name"].split(".")[-1],
                        "extracted_text": att.get("extracted_text") or ""
                    })
                break

    remote_response = ""
    for msg in reversed(db_messages):
        if msg["role"] == "assistant" and msg["content"] and "Hi, by looking" in msg["content"]:
            remote_response = msg["content"]
            break
    if not remote_response:
        for msg in reversed(db_messages):
            if msg["role"] == "assistant" and msg["content"] and msg["id"] != user_msg_id:
                remote_response = msg["content"]
                break

    # 2. Setup progress queue and run local document processing workflow
    queue = asyncio.Queue()
    loop = asyncio.get_running_loop()
    all_progress_events = []
    
    def on_progress(event_type: str, data: dict):
        event = {"type": event_type, **data}
        all_progress_events.append(event)
        loop.call_soon_threadsafe(queue.put_nowait, event)

    from app.services.agent_workflow import local_document_processing_workflow
    
    input_data = {
        "remote_response": remote_response,
        "files_list": files_list,
        "session_id": session_id,
        "on_progress": on_progress,
        "user_feedback": last_user_message
    }

    async def local_chat_generator():
        try:
            # Yield session_created event
            yield f"data: {json.dumps({'type': 'session_created', 'session_id': session_id})}\n\n"

            # Start workflow task
            workflow_task = asyncio.create_task(local_document_processing_workflow.run(input_data))

            # Stream progress events
            while not workflow_task.done() or not queue.empty():
                try:
                    event = await asyncio.wait_for(queue.get(), timeout=0.1)
                    yield f"data: {json.dumps(event)}\n\n"
                except asyncio.TimeoutError:
                    continue

            # Await final result
            run_result = await workflow_task
            outputs = run_result.get_outputs()
            result_dict = outputs[-1] if outputs else {}

            if isinstance(result_dict, dict):
                final_text = result_dict.get("final_response", "Failed to process.")
                processed_files = result_dict.get("processed_files", [])
                citation_map = result_dict.get("citation_map", {})
            else:
                final_text = result_dict if isinstance(result_dict, str) else "Failed to process."
                processed_files = []
                citation_map = {}

            # 3. Save assistant response to DB
            assistant_msg_id = str(uuid.uuid4())
            save_message(
                assistant_msg_id,
                session_id,
                "assistant",
                final_text,
                datetime.utcnow().isoformat(),
                progress_events=json.dumps(all_progress_events),
                citation_map=json.dumps(citation_map) if citation_map else None
            )

            # 4. Save processed attachments in DB linked to assistant response
            upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "uploads", session_id)
            for pf in processed_files:
                fname = pf.get("filename")
                if not fname:
                    continue
                fpath = os.path.join(upload_dir, fname)
                if os.path.exists(fpath) and os.path.isfile(fpath):
                    save_attachment(
                        assistant_msg_id,
                        fname,
                        os.path.getsize(fpath),
                        f"/uploads/{session_id}/{fname}",
                        extracted_text=pf.get("extracted_text", "")
                    )

            # Emit done event with payload
            yield f"data: {json.dumps({'type': 'done', 'final_response': final_text, 'processed_files': processed_files, 'citation_map': citation_map})}\n\n"

        except Exception as e:
            logger.error(f"Error handling local follow-up chat workflow: {e}")
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return StreamingResponse(local_chat_generator(), media_type="text/event-stream")


@router.post("/chat")
async def handle_chat(payload: ChatRequest):
    # Ensure Azure client is initialized
    try:
        azure_client_manager.initialize()
    except Exception as e:
        logger.error(f"Failed to initialize Azure AI Client: {e}")
        raise HTTPException(status_code=500, detail=f"Azure Client error: {str(e)}")
        
    try:
        openai_client = azure_client_manager.get_openai_client()
        
        # 1. Determine or generate session ID
        session_id = payload.session_id
        is_new_session = False
        if not session_id:
            session_id = str(uuid.uuid4())
            is_new_session = True

        # Determine user query
        last_user_message = ""
        if payload.messages:
            for m in reversed(payload.messages):
                if m.role == "user" and m.content:
                    last_user_message = m.content
                    break
        if not last_user_message:
            last_user_message = payload.message or ""

        # Determine user image (if any)
        last_user_image = None
        if payload.messages:
            for m in reversed(payload.messages):
                if m.role == "user" and m.image:
                    last_user_image = m.image
                    break
        if not last_user_image:
            last_user_image = payload.image

        # If it's a new session or we need to save the new user message
        if is_new_session:
            # Generate title from message or default
            title = last_user_message.strip()[:30] + "..." if len(last_user_message.strip()) > 30 else (last_user_message.strip() or "Form Analysis")
            create_session(session_id, title)

        # Save user message to database
        user_msg_id = str(uuid.uuid4())
        save_message(
            user_msg_id,
            session_id,
            "user",
            last_user_message,
            datetime.utcnow().isoformat()
        )

        # Save user attachment (if image present)
        if last_user_image:
            save_attachment(
                user_msg_id,
                "Screenshot.png",
                None,
                last_user_image
            )

        # 2. Build the input messages array from SQLite history
        db_messages = get_session_messages(session_id)
        input_messages = []
        for msg in db_messages:
            if msg["role"] == "assistant":
                item = {"role": "assistant"}
                if msg["content"]:
                    item["content"] = msg["content"]
                input_messages.append(item)
            else:
                # user message content
                attachments = msg.get("attachments", [])
                image_data_url = None
                for att in attachments:
                    if att["data_url"] and att.get("data_url").startswith("data:image/"):
                        image_data_url = att["data_url"]
                        break
                
                if image_data_url:
                    content = [
                        {"type": "input_text", "text": msg["content"] or ""},
                        {"type": "input_image", "image_url": image_data_url}
                    ]
                else:
                    content = msg["content"] or ""
                    
                input_messages.append({
                    "role": "user",
                    "content": content
                })

        logger.info(f"Streaming chat request to remote responses API with {len(input_messages)} messages")
        
        def remote_event_generator():
            try:
                # Yield session_created event first
                yield f"data: {json.dumps({'type': 'session_created', 'session_id': session_id})}\n\n"

                # Call responses API with stream=True
                response_stream = openai_client.responses.create(
                    input=input_messages,
                    stream=True
                )
                
                assistant_text = ""
                all_annotations = []
                for chunk in response_stream:
                    # 1. Output text delta chunk
                    if chunk.type == "response.output_text.delta":
                        assistant_text += chunk.delta
                        data = {"type": "text_delta", "text": chunk.delta}
                        yield f"data: {json.dumps(data)}\n\n"
                    
                    # 2. Annotation added event
                    elif chunk.type == "response.output_text.annotation.added":
                        ann = chunk.annotation
                        annotation_dict = {}
                        if hasattr(ann, "model_dump"):
                          annotation_dict = ann.model_dump()
                        elif hasattr(ann, "dict"):
                          annotation_dict = ann.dict()
                        elif hasattr(ann, "__dict__"):
                          annotation_dict = {k: v for k, v in ann.__dict__.items() if not k.startswith('_')}
                        else:
                          annotation_dict = dict(ann)
                            
                        # Add chunk level properties for marker matching
                        annotation_dict["output_index"] = getattr(chunk, "output_index", 0)
                        annotation_dict["annotation_index"] = getattr(chunk, "annotation_index", 0)
                        
                        all_annotations.append(annotation_dict)
                        data = {"type": "annotation", "annotation": annotation_dict}
                        yield f"data: {json.dumps(data)}\n\n"
                    
                    # 3. Done event
                    elif chunk.type == "response.done":
                        break
 
                # Save assistant message to SQLite
                save_message(
                    str(uuid.uuid4()),
                    session_id,
                    "assistant",
                    assistant_text,
                    datetime.utcnow().isoformat(),
                    annotations=json.dumps(all_annotations) if all_annotations else None
                )
                yield f"data: {json.dumps({'type': 'done'})}\n\n"
            except Exception as stream_err:
                logger.error(f"Error in remote event stream generation: {stream_err}")
                yield f"data: {json.dumps({'type': 'error', 'message': str(stream_err)})}\n\n"

        return StreamingResponse(remote_event_generator(), media_type="text/event-stream")
        
    except Exception as e:
        logger.error(f"Error calling responses API stream: {e}")
        raise HTTPException(status_code=500, detail=f"Agent response stream error: {str(e)}")
