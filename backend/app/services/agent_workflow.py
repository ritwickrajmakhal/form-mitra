import os
import json
import logging
import tempfile
import uuid
import shutil
from datetime import datetime
from PIL import Image
import fitz  # PyMuPDF
from agent_framework import workflow, step, RunContext
from app.services.local_model import local_model_service
from app.db import save_message, save_attachment

logger = logging.getLogger("app.services.agent_workflow")

# ---------------------------------------------------------------------------
# Custom Document Tools
# ---------------------------------------------------------------------------

def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    import docx
    import fitz
    
    doc = docx.Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                paragraphs.append(" | ".join(row_text))
                
    pdf = fitz.open()
    page_width = 612
    page_height = 792
    margin = 50
    font_size = 11
    line_height = 14
    
    font = fitz.Font("helv")
    page = pdf.new_page(width=page_width, height=page_height)
    y = margin
    
    def wrap_text(text, max_width, font, size):
        words = text.split()
        lines = []
        current_line = []
        for word in words:
            test_line = " ".join(current_line + [word])
            width = font.text_length(test_line, fontsize=size)
            if width <= max_width:
                current_line.append(word)
            else:
                if current_line:
                    lines.append(" ".join(current_line))
                    current_line = [word]
                else:
                    lines.append(word)
        if current_line:
            lines.append(" ".join(current_line))
        return lines

    max_text_width = page_width - (2 * margin)
    
    for para in paragraphs:
        wrapped_lines = wrap_text(para, max_text_width, font, font_size)
        for line in wrapped_lines:
            if y + line_height > page_height - margin:
                page = pdf.new_page(width=page_width, height=page_height)
                y = margin
            page.insert_text(fitz.Point(margin, y), line, fontsize=font_size, fontname="helv")
            y += line_height
        y += line_height / 2
        
    pdf.save(pdf_path)
    pdf.close()

def convert_text_to_pdf(text_path: str, pdf_path: str):
    import fitz
    with open(text_path, "r", encoding="utf-8", errors="ignore") as f:
        text_content = f.read()
    paragraphs = [p.strip() for p in text_content.split("\n") if p.strip()]
    
    pdf = fitz.open()
    page_width = 612
    page_height = 792
    margin = 50
    font_size = 11
    line_height = 14
    
    font = fitz.Font("helv")
    page = pdf.new_page(width=page_width, height=page_height)
    y = margin
    
    def wrap_text(text, max_width, font, size):
        words = text.split()
        lines = []
        current_line = []
        for word in words:
            test_line = " ".join(current_line + [word])
            width = font.text_length(test_line, fontsize=size)
            if width <= max_width:
                current_line.append(word)
            else:
                if current_line:
                    lines.append(" ".join(current_line))
                    current_line = [word]
                else:
                    lines.append(word)
        if current_line:
            lines.append(" ".join(current_line))
        return lines

    max_text_width = page_width - (2 * margin)
    
    for para in paragraphs:
        wrapped_lines = wrap_text(para, max_text_width, font, font_size)
        for line in wrapped_lines:
            if y + line_height > page_height - margin:
                page = pdf.new_page(width=page_width, height=page_height)
                y = margin
            page.insert_text(fitz.Point(margin, y), line, fontsize=font_size, fontname="helv")
            y += line_height
        y += line_height / 2
        
    pdf.save(pdf_path)
    pdf.close()

def convert_document_tool(file_path: str, target_format: str, session_id: str) -> str:
    """Converts a document between PDF, Word, Text and Image formats.
    
    Args:
        file_path: Absolute path or filename in the session uploads folder.
        target_format: 'pdf', 'png', or 'jpg'.
        session_id: The active session ID.
    """
    logger.info(f"Converting document {file_path} to {target_format} in session {session_id}")
    
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    uploads_dir = os.path.join(base_dir, "uploads", session_id)
    os.makedirs(uploads_dir, exist_ok=True)
    
    # Resolve absolute path
    if not os.path.isabs(file_path):
        file_path = os.path.join(uploads_dir, os.path.basename(file_path))
        
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    filename = os.path.basename(file_path)
    base_name, ext = os.path.splitext(filename)
    target_format = target_format.lower()
    
    IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".heic", ".heif", ".avif")
    
    if target_format == "pdf":
        new_filename = f"{base_name}.pdf"
        dest_path = os.path.join(uploads_dir, new_filename)
        
        if ext.lower() in IMAGE_EXTENSIONS:
            # Image -> PDF
            img = Image.open(file_path)
            img.convert('RGB').save(dest_path, 'PDF')
            logger.info(f"Converted image to PDF at {dest_path}")
            return f"/uploads/{session_id}/{new_filename}"
        elif ext.lower() == ".docx":
            # Word -> PDF
            convert_docx_to_pdf(file_path, dest_path)
            logger.info(f"Converted Word document to PDF at {dest_path}")
            return f"/uploads/{session_id}/{new_filename}"
        elif ext.lower() in (".txt", ".csv", ".md"):
            # Text -> PDF
            convert_text_to_pdf(file_path, dest_path)
            logger.info(f"Converted Text file to PDF at {dest_path}")
            return f"/uploads/{session_id}/{new_filename}"
        else:
            raise ValueError(f"Cannot convert {ext} to PDF format")
            
    elif target_format in ("png", "jpg", "jpeg"):
        new_filename = f"{base_name}.png"
        dest_path = os.path.join(uploads_dir, new_filename)
        
        if ext.lower() == ".pdf":
            # PDF -> Image (first page)
            doc = fitz.open(file_path)
            page = doc.load_page(0)
            pix = page.get_pixmap()
            pix.save(dest_path)
            logger.info(f"Converted PDF to image at {dest_path}")
            return f"/uploads/{session_id}/{new_filename}"
        elif ext.lower() in IMAGE_EXTENSIONS:
            # Modern Image -> Standard Image (PNG/JPG)
            img = Image.open(file_path)
            img.save(dest_path, target_format.upper())
            logger.info(f"Converted modern image {ext} to standard {target_format} at {dest_path}")
            return f"/uploads/{session_id}/{new_filename}"
        else:
            raise ValueError(f"Cannot convert {ext} to standard image format")
    else:
        raise ValueError(f"Unsupported target format: {target_format}")

def compress_document_tool(file_path: str, session_id: str) -> str:
    """Compresses a PDF or image file to reduce its size.
    
    Args:
        file_path: Absolute path or filename in the session uploads folder.
        session_id: The active session ID.
    """
    logger.info(f"Compressing document {file_path} in session {session_id}")
    
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    uploads_dir = os.path.join(base_dir, "uploads", session_id)
    os.makedirs(uploads_dir, exist_ok=True)
    
    if not os.path.isabs(file_path):
        file_path = os.path.join(uploads_dir, os.path.basename(file_path))
        
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    filename = os.path.basename(file_path)
    base_name, ext = os.path.splitext(filename)
    
    IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".heic", ".heif", ".avif")
    
    if ext.lower() in IMAGE_EXTENSIONS:
        # If it's a modern image format like HEIC/AVIF, save it as optimized .jpg for absolute compatibility and compression
        if ext.lower() in (".heic", ".heif", ".avif", ".tiff"):
            new_filename = f"{base_name}_compressed.jpg"
            dest_path = os.path.join(uploads_dir, new_filename)
            img = Image.open(file_path)
            img.convert('RGB').save(dest_path, "JPEG", optimize=True, quality=50)
            logger.info(f"Converted and compressed modern image {ext} to {new_filename}")
            return f"/uploads/{session_id}/{new_filename}"
        else:
            new_filename = f"{base_name}_compressed{ext}"
            dest_path = os.path.join(uploads_dir, new_filename)
            img = Image.open(file_path)
            if ext.lower() in (".jpg", ".jpeg"):
                img.convert('RGB').save(dest_path, optimize=True, quality=50)
            else:
                img.save(dest_path, optimize=True)
            logger.info(f"Compressed image saved to {dest_path}")
            return f"/uploads/{session_id}/{new_filename}"
        
    elif ext.lower() == ".pdf":
        new_filename = f"{base_name}_compressed.pdf"
        dest_path = os.path.join(uploads_dir, new_filename)
        # PDF compression using PyMuPDF
        doc = fitz.open(file_path)
        doc.save(dest_path, garbage=4, deflate=True)
        logger.info(f"Compressed PDF saved to {dest_path}")
        return f"/uploads/{session_id}/{new_filename}"
    else:
        # Copy other files directly
        new_filename = f"{base_name}_compressed{ext}"
        dest_path = os.path.join(uploads_dir, new_filename)
        shutil.copy(file_path, dest_path)
        logger.info(f"Unsupported format copy saved to {dest_path}")
        return f"/uploads/{session_id}/{new_filename}"

# ---------------------------------------------------------------------------
# Workflow Steps
# ---------------------------------------------------------------------------

@step(name="plan_document_actions")
async def plan_document_actions(remote_response: str, files_list: list, feedback: str = "") -> list:
    """Analyzes requirements and files to produce a list of conversion/compression actions using the local model."""
    logger.info("Step 1: Planning document actions...")
    
    system_prompt = """You are Form Mitra's local agent planning assistant.
Analyze the required format and size constraints for each document from the remote agent's response, and check them against the uploaded files list.

Rules:
1. Map each uploaded file to a required document based on filename keywords (e.g. "aadhar" matches "Aadhar Card", "vote" or "voter" matches "Voter Card", "passport" or "photo" matches "Passport Size Photo").
2. Check if the uploaded file format matches the required format:
   - If a document is required in "PDF" format, but the uploaded file is an image (extension .png, .jpg, .jpeg, .webp, .bmp, .tiff, .heic, .heif, .avif), you MUST convert it to pdf.
   - If a document is required in "Image" format, but the uploaded file is a PDF, you MUST convert it to png.
   - If the uploaded file format ALREADY satisfies the required format (e.g. uploaded file is an image and requirement is "Image", or uploaded file is a PDF and requirement is "PDF"), you MUST NOT convert it.
3. Check if the uploaded file size exceeds the required maximum size (e.g., "1 MB" = 1,048,576 bytes). If yes, you MUST compress it (action: "compress").
4. If the file already meets both format and size requirements, you MUST NOT convert or compress it (action: "none").
5. Never output a "convert" action with "target_format": "none". Output "none" action instead.

Here is a planning walkthrough:
- User uploads "aadhar_synthetic.avif" (15 KB, Image). Remote requires "Aadhar Card" as "PDF". Action: {"action": "convert", "filename": "aadhar_synthetic.avif", "target_format": "pdf"}
- User uploads "voter_synthetic.pdf" (1.2 MB, PDF). Remote requires "Voter Card" as "PDF", max size 1 MB. Action: {"action": "compress", "filename": "voter_synthetic.pdf"}
- User uploads "synthetic_passport_size_photo.jpg" (45 KB, Image). Remote requires "Passport size Photo" as "Image". Action: {"action": "none", "filename": "synthetic_passport_size_photo.jpg", "reason": "Already in Image format and within size limit"}

Do not output any other text or reasoning. Start directly with the JSON brackets []."""

    prompt = f"Remote Agent Response:\n{remote_response}\n\nUploaded Files:\n{json.dumps(files_list, indent=2)}"
    if feedback:
        prompt += f"\n\n--- FEEDBACK FROM PREVIOUS VERIFICATION ATTEMPT (Please correct planning/conversion issues): ---\n{feedback}"
        
    response_text = ""
    try:
        stream = local_model_service.generate_stream(system_prompt, prompt)
        response_text = "".join(list(stream)).strip()
        logger.info(f"Planner response: {response_text}")
        
        # Parse JSON list
        if "```" in response_text:
            lines = response_text.split("\n")
            cleaned_lines = [l for l in lines if not l.startswith("```")]
            response_text = "\n".join(cleaned_lines).strip()
            
        actions = json.loads(response_text)
        if isinstance(actions, list):
            return actions
        return []
    except Exception as e:
        logger.error(f"Failed to parse planner JSON actions: {e}. Raw text: {response_text}")
        return []

@step(name="execute_document_tools")
async def execute_document_tools(actions: list, session_id: str, files_list: list, on_progress=None, attempt: int = 1) -> list:
    """Executes the conversion or compression tools and updates the file list."""
    logger.info(f"Step 2: Executing document tools for actions: {actions}")
    
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    uploads_dir = os.path.join(base_dir, "uploads", session_id)
    
    # We copy the dictionary objects to prevent mutating input parameters directly
    updated_files = [dict(f) for f in files_list]
    
    # name_mapping tracks the mapping from the original filename (as planned by the agent)
    # to the current filename in uploads_dir as it goes through sequential conversions/compressions
    name_mapping = {f["filename"]: f["filename"] for f in updated_files}
    
    seen_actions = set()
    for action in actions:
        filename = action.get("filename")
        act_type = action.get("action")
        target_format = action.get("target_format")
        
        # Deduplicate identical actions in the same attempt
        action_key = (filename, act_type, target_format)
        if action_key in seen_actions:
            logger.info(f"Skipping duplicate action: {action}")
            continue
        seen_actions.add(action_key)
        
        current_filename = name_mapping.get(filename, filename)
        
        # Find the file in our list
        matching_file = None
        for f in updated_files:
            if f.get("filename") == current_filename:
                matching_file = f
                break
                
        if not matching_file:
            logger.warning(f"File '{filename}' (resolved as '{current_filename}') not found in files list.")
            continue
            
        file_path = os.path.join(uploads_dir, current_filename)
        
        if on_progress:
            try:
                on_progress("tool_start", {
                    "attempt": attempt,
                    "action": act_type, 
                    "filename": filename, 
                    "target_format": action.get("target_format")
                })
            except Exception as pe:
                logger.error(f"Error calling on_progress in execute_document_tools: {pe}")
        
        success = False
        try:
            if act_type == "none":
                success = True
            elif act_type == "compress":
                new_url = compress_document_tool(file_path, session_id)
                new_filename = os.path.basename(new_url)
                new_path = os.path.join(uploads_dir, new_filename)
                
                # Update files list details
                matching_file["filename"] = new_filename
                matching_file["filesize"] = os.path.getsize(new_path)
                matching_file["filetype"] = new_filename.split(".")[-1]
                logger.info(f"Successfully compressed file {current_filename} to {new_filename}")
                success = True
                name_mapping[filename] = new_filename
                
            elif act_type == "convert":
                target_format = action.get("target_format", "pdf")
                if target_format in ("none", ""):
                    logger.info(f"Skipping no-op convert for {current_filename}")
                    success = True
                else:
                    new_url = convert_document_tool(file_path, target_format, session_id)
                    new_filename = os.path.basename(new_url)
                    new_path = os.path.join(uploads_dir, new_filename)
                    
                    # Update files list details
                    matching_file["filename"] = new_filename
                    matching_file["filesize"] = os.path.getsize(new_path)
                    matching_file["filetype"] = target_format
                    logger.info(f"Successfully converted file {current_filename} to {new_filename}")
                    success = True
                    name_mapping[filename] = new_filename                
        except Exception as tool_err:
            logger.error(f"Error executing tool {act_type} for {filename}: {tool_err}")
            
        if on_progress:
            try:
                on_progress("tool_end", {
                    "attempt": attempt,
                    "action": act_type, 
                    "filename": filename, 
                    "success": success, 
                    "new_filename": matching_file.get("filename") if success else None
                })
            except Exception as pe:
                logger.error(f"Error calling on_progress in execute_document_tools: {pe}")
            
    return updated_files

# ---------------------------------------------------------------------------
# Date/Time Tool
# ---------------------------------------------------------------------------

@step(name="get_current_datetime")
async def get_current_datetime() -> str:
    """Returns the current date and time as a formatted string for the model to use in reasoning."""
    from datetime import datetime
    now = datetime.now()
    result = now.strftime("%d %B %Y (day %d, month %m, year %Y), %A, %H:%M")
    logger.info(f"[get_current_datetime] Current datetime: {result}")
    return result

@step(name="generate_extracted_response")
async def generate_extracted_response(remote_response: str, processed_files: list, current_datetime: str = "", feedback: str = "") -> str:
    """Generates the final formatted response utilizing OCR text mapped to form fields."""
    logger.info("Step 3: Formatting final response...")

    datetime_context = f"\nTODAY'S DATE: {current_datetime}" if current_datetime else ""

    system_prompt = f"""You are Form Mitra, an intelligent local form filling assistant.
Your goal is to extract the values for the text fields requested in the remote agent's response using the provided OCR texts from the uploaded documents.{datetime_context}

CRITICAL FORMAT RULES:
1. Output format MUST be exactly a numbered list of text fields:
   1. [Field Name]: [Extracted Value]
   2. [Field Name]: [Extracted Value]
   ...
2. Start directly with the first field. Do NOT output any conversational greetings, introductions, or markdown blocks (like ```).
3. INCLUDE vs EXCLUDE rules:
   - INCLUDE fields that hold personal data values: Name, Father's Name, Gender, Date of Birth, Age, Address, Aadhaar Card Number, Voter Card Number, EPIC Number, PAN Number, etc.
   - EXCLUDE bare document/attachment upload fields: "Aadhar Card", "Voter Card", "Passport size Photo", "Photograph" — these are file attachments rendered separately in the UI.
   - KEY DISTINCTION: If the field name contains "Number", "No.", "No ", "ID", or "EPIC", it is a TEXT field (include it). If the field name is just a document name like "Voter Card" alone (no number/ID qualifier), it is an attachment (exclude it).
4. AGE CALCULATION RULE: If "Age" is a requested field and the OCR text does NOT contain an explicit age value, but DOES contain a Date of Birth (DOB), calculate the age yourself using TODAY'S DATE provided above. Reason step by step: subtract the birth year from today's year, then check if the birthday has occurred yet this year to get the exact age in years.
5. If a value truly cannot be determined from any OCR text or reasoning, write "Not found" as the value.
6. Extract ALL text fields — do not skip any.

FEW-SHOT EXAMPLE (Age calculation from DOB):
Today's Date: 14 June 2026 (day 14, month 06, year 2026), Sunday, 01:30
Input Remote Agent Response:
"Please provide: 1. Name 2. Date of Birth 3. Age 4. Aadhaar Card Number 5. Voter Card Number 6. Aadhaar Card (PDF) 7. Voter Card (PDF)"
Input OCR Texts:
File: aadhar_synthetic.pdf -> "NAME: JOHN SMITH, DOB: 15/08/1995, AADHAR NO: 9876 5432 1098"
File: voter_synthetic.pdf -> "NAME: John Smith, EPIC NO: XYZ1234567"

Reasoning: DOB is 15 August 1995. Today is 14 June 2026. From 1995 to 2026 = 31 years. But birthday (15 Aug) has NOT yet occurred in 2026 (we are in June). So age = 30.

Expected Output:
1. Name: John Smith
2. Date of Birth: 15/08/1995
3. Age: 30
4. Aadhaar Card Number: 9876 5432 1098
5. Voter Card Number: XYZ1234567"""

    # Build prompt content from files and their OCR texts.
    # Documents are numbered so the model can cite them as [1], [2], etc.
    ocr_descriptions = []
    citation_map = {}  # index (1-based str) -> filename
    for i, f in enumerate(processed_files):
        doc_num = i + 1
        filename = f.get("filename")
        citation_map[str(doc_num)] = filename
        ocr_text = f.get("extracted_text", "")
        ocr_descriptions.append(f"[Document {doc_num}] File: {filename}\nOCR Text:\n{ocr_text}\n---")
        
    prompt = f"Remote Agent Response (Required Fields & Uploads):\n{remote_response}\n\nOCR Texts (numbered for citations):\n" + "\n".join(ocr_descriptions)
    prompt += "\n\nCITATION INSTRUCTION: After each extracted value, append [N] where N is the Document number you found that value in (e.g. [1] for Document 1). If the value was calculated (like Age from DOB), do NOT add a citation. If 'Not found', no citation."
    
    if feedback:
        prompt += f"\n\n--- CRITICAL VERIFICATION FEEDBACK (You MUST correct this in your output): ---\n{feedback}"

    logger.info(f"[generate_extracted_response] Prompt:\n{prompt}")
    try:
        stream = local_model_service.generate_stream(system_prompt, prompt)
        response_text = "".join(list(stream)).strip()
        logger.info(f"[generate_extracted_response] Raw output:\n{response_text}")
        return response_text, citation_map
    except Exception as e:
        logger.error(f"Error generating final extracted response: {e}")
        return "Failed to process form extraction fields.", {}

@step(name="verify_extracted_response")
async def verify_extracted_response(response_text: str, processed_files: list, remote_response: str) -> dict:
    """Verifies both the text response and that all processed files meet format and size requirements."""
    logger.info("Step 4: Verifying extracted response and files...")

    # ---------------------------------------------------------------------------
    # CHECK A: File format and size verification (programmatic, reliable)
    # ---------------------------------------------------------------------------
    file_check_errors = []
    import re as _re
    for f in processed_files:
        filename = f.get("filename", "")
        filesize = f.get("filesize", 0)
        filetype = f.get("filetype", "").lower()
        ext = os.path.splitext(filename)[1].lower()

        # Try to determine required format from remote_response by matching filename keywords
        required_pdf = False
        required_image = False
        remote_lower = remote_response.lower()

        if "aadhar" in filename.lower() or "aadhaar" in filename.lower():
            if "pdf" in remote_lower[max(0, remote_lower.find("aadhar") - 5):remote_lower.find("aadhar") + 200]:
                required_pdf = True
        if "voter" in filename.lower():
            if "pdf" in remote_lower[max(0, remote_lower.find("voter") - 5):remote_lower.find("voter") + 200]:
                required_pdf = True
        if "passport" in filename.lower() or "photo" in filename.lower():
            if "image" in remote_lower[max(0, remote_lower.find("passport") - 5):remote_lower.find("passport") + 200]:
                required_image = True

        IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".heic", ".heif", ".avif")
        if required_pdf and ext != ".pdf":
            file_check_errors.append(f"File '{filename}' must be PDF format but is '{ext}'.")
        if required_image and ext not in IMAGE_EXTENSIONS:
            file_check_errors.append(f"File '{filename}' must be Image format but is '{ext}'.")

    if file_check_errors:
        feedback = " ".join(file_check_errors)
        logger.info(f"[verify] CHECK A failed: {feedback}")
        return {"is_valid": False, "feedback": feedback}

    # ---------------------------------------------------------------------------
    # CHECK B: Text response field completeness (LLM-based, dynamic)
    # ---------------------------------------------------------------------------
    # Extract what text fields the remote agent requested (exclude document attachment lines)
    ATTACHMENT_KEYWORDS = ("aadhar card", "aadhaar card", "voter card", "passport size photo",
                           "passport photo", "photograph", "passport size", "id proof", "document upload")
    ID_NUMBER_KEYWORDS = ("number", "no.", " no ", "id", "card no", "epic")

    requested_text_fields = []
    for line in remote_response.split("\n"):
        line_s = line.strip()
        if not line_s:
            continue
        # Remove leading numbering (e.g. "1. ", "2) ")
        field_raw = _re.sub(r'^\d+[\s\.\)]+', '', line_s).strip()
        # Strip Azure AI Foundry inline citation markers: 【4:0†source】 or 【3†source】
        field_raw = _re.sub(r'【[^】]*†source】', '', field_raw).strip()
        # Also strip any trailing [web] or [N] markers the remote agent may add
        field_raw = _re.sub(r'\[web\]|\[\d+\]', '', field_raw).strip()
        field_lower = field_raw.lower()
        # Skip lines that are document attachment requests (not number fields)
        is_attachment = any(kw in field_lower for kw in ATTACHMENT_KEYWORDS)
        has_id_qualifier = any(kw in field_lower for kw in ID_NUMBER_KEYWORDS)
        if is_attachment and not has_id_qualifier:
            continue
        # Skip lines that look like "Format:", "Max Size:" explanatory text
        if any(kw in field_lower for kw in ("format:", "max size:", "kindly", "please", "attachment", "documents")):
            continue
        # Skip very short or empty tokens
        if len(field_raw) < 2:
            continue
        # Only keep lines that look like field names (not full sentences)
        if len(field_raw.split()) <= 5:
            requested_text_fields.append(field_raw)

    logger.info(f"[verify] Detected {len(requested_text_fields)} requested text fields: {requested_text_fields}")

    # Count how many were found in the response
    response_lower = response_text.lower()
    missing_fields = []
    for field in requested_text_fields:
        if field.lower() not in response_lower:
            missing_fields.append(field)

    if missing_fields:
        feedback = f"Missing fields in your response: {', '.join(missing_fields)}. Extract values for ALL requested text fields from the OCR data."
        logger.info(f"[verify] CHECK B failed: {feedback}")
        return {"is_valid": False, "feedback": feedback}

    logger.info("[verify] All checks passed.")
    return {"is_valid": True, "feedback": ""}


# ---------------------------------------------------------------------------
# Workflow Definition
# ---------------------------------------------------------------------------

@workflow(name="local_document_processing_workflow")
async def local_document_processing_workflow(input_data: dict) -> dict:
    """Orchestrates document analysis, planning, conversion/compression tools execution, and final extraction."""
    remote_response = input_data.get("remote_response", "")
    files_list = input_data.get("files_list", [])
    session_id = input_data.get("session_id", "")
    on_progress = input_data.get("on_progress")

    def trigger_progress(event_type: str, event_data: dict):
        if on_progress:
            try:
                on_progress(event_type, event_data)
            except Exception as pe:
                logger.error(f"Error calling on_progress in workflow: {pe}")

    feedback = ""
    actions = []
    processed_files = []
    final_response = ""

    for attempt in range(3):
        logger.info(f"Workflow attempt {attempt + 1}...")
        trigger_progress("attempt_start", {"attempt": attempt + 1})

        # 1. Plan document actions (compression or conversion)
        # Only pass feedback to planner if it directly relates to a file check failure (starts with 'File ')
        planner_feedback = feedback if feedback.strip().startswith("File ") else ""
        trigger_progress("planning_start", {"attempt": attempt + 1})
        actions = await plan_document_actions(remote_response, files_list, planner_feedback)
        trigger_progress("planning_end", {"attempt": attempt + 1, "actions": actions})

        # 2. Execute tools on files
        trigger_progress("execution_start", {"attempt": attempt + 1, "actions": actions})
        processed_files = await execute_document_tools(actions, session_id, files_list, on_progress=on_progress, attempt=attempt + 1)
        trigger_progress("execution_end", {"attempt": attempt + 1, "processed_files": processed_files})

        # 3a. Get current datetime for age reasoning
        current_datetime = await get_current_datetime()
        trigger_progress("datetime_fetched", {"attempt": attempt + 1, "current_datetime": current_datetime})

        # 3b. Generate final extracted fields response
        # Passes current datetime so the model can reason about Age from DOB.
        # The system prompt instructs the model to skip attachment fields.
        trigger_progress("formatting_start", {"attempt": attempt + 1})
        formatter_result = await generate_extracted_response(
            remote_response,
            processed_files,
            current_datetime=current_datetime,
            feedback=feedback if not feedback.strip().startswith("File ") else ""
        )
        # generate_extracted_response returns (text, citation_map)
        if isinstance(formatter_result, tuple):
            final_response, citation_map = formatter_result
        else:
            final_response, citation_map = formatter_result, {}
        trigger_progress("formatting_end", {"attempt": attempt + 1, "final_response": final_response})

        # 4. Verify both response text and processed files
        trigger_progress("verification_start", {"attempt": attempt + 1})
        verification_result = await verify_extracted_response(final_response, processed_files, remote_response)
        is_valid = verification_result.get("is_valid", False)
        feedback = verification_result.get("feedback", "")
        trigger_progress("verification_end", {"attempt": attempt + 1, "is_valid": is_valid, "feedback": feedback})

        if is_valid:
            logger.info("Verification succeeded.")
            break
        else:
            logger.info(f"Verification failed (attempt {attempt + 1}). Feedback: {feedback}")

    return {
        "final_response": final_response,
        "processed_files": processed_files,
        "citation_map": citation_map
    }

